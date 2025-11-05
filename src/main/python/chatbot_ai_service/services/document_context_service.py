"""
Servicio para cargar y gestionar documentos del cliente usando LlamaIndex

Este servicio carga documentos desde el bucket del cliente y los indexa
para proporcionar contexto relevante a la IA.
"""
import logging
import os
import tempfile
from typing import Dict, Any, Optional, List
from urllib.parse import urlparse
import httpx
import io

# Intentar importar LlamaIndex (rutas varían según versión). Si falla, se desactiva RAG.
VectorStoreIndex = None
Document = None
Settings = None
StorageContext = None
load_index_from_storage = None
SimpleNodeParser = None
Gemini = None
GeminiEmbedding = None
import google.generativeai as genai

# Importar librerías para procesamiento de documentos
try:
    import pypdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("pypdf no disponible - soporte PDF deshabilitado")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx no disponible - soporte Word deshabilitado")

try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False
    logging.warning("openpyxl no disponible - soporte Excel deshabilitado")

# Verificar que LlamaIndex esté disponible y resolver rutas según versión
LLAMA_INDEX_SUPPORT = False
try:
    # Rutas modernas (>=0.9.x)
    from llama_index.core import (
        VectorStoreIndex as _VSI,
        Document as _Doc,
        Settings as _Settings,
        StorageContext as _SC,
        load_index_from_storage as _load
    )
    from llama_index.core.node_parser import SimpleNodeParser as _SNP
    try:
        from llama_index.llms.gemini import Gemini as _Gem
        from llama_index.embeddings.gemini import GeminiEmbedding as _GemEmb
    except ImportError as e:
        logging.warning(f"⚠️ No se pudo importar Gemini para LlamaIndex: {str(e)}")
        _Gem = None
        _GemEmb = None
    VectorStoreIndex, Document, Settings, StorageContext, load_index_from_storage = _VSI, _Doc, _Settings, _SC, _load
    SimpleNodeParser, Gemini, GeminiEmbedding = _SNP, _Gem, _GemEmb
    LLAMA_INDEX_SUPPORT = True
    logging.info("✅ LlamaIndex cargado correctamente (modo moderno)")
except ImportError as e:
    logging.error(f"❌ Error cargando LlamaIndex moderno: {str(e)}")
    try:
        # Rutas legacy (<0.9)
        from llama_index import (
            VectorStoreIndex as _VSI,
            Document as _Doc,
            ServiceContext as _Settings
        )
        from llama_index.node_parser import SimpleNodeParser as _SNP
        VectorStoreIndex, Document, Settings = _VSI, _Doc, _Settings
        SimpleNodeParser = _SNP
        LLAMA_INDEX_SUPPORT = True
        logging.warning("⚠️ LlamaIndex en modo legacy (sin Gemini integrado)")
    except ImportError as e2:
        logging.error(f"❌ LlamaIndex no disponible: {str(e2)}")
        logging.warning("⚠️ Documentos se cargarán sin indexación vectorial (modo simple)")

logger = logging.getLogger(__name__)

class DocumentContextService:
    """Servicio para cargar documentos del cliente y proporcionar contexto a la IA"""
    
    def __init__(self):
        self._index_cache = {}  # Cache de índices por tenant
        self._document_cache = {}  # Cache de documentos cargados
        self._llm = None
        self._embedding_model = None
        self._initialized = False
    
    def _ensure_models_initialized(self):
        """Inicializa los modelos de IA de forma lazy"""
        if self._initialized:
            return
            
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key and LLAMA_INDEX_SUPPORT and Gemini is not None:
            try:
                genai.configure(api_key=api_key)
                
                # Configurar modelo por defecto de LlamaIndex primero
                if Settings is not None:
                    # Crear modelo por defecto para evitar errores
                    default_llm = Gemini(model="gemini-2.0-flash-exp", api_key=api_key)
                    default_embedding = GeminiEmbedding(
                        model_name="models/embedding-001",
                        api_key=api_key
                    )
                    Settings.llm = default_llm
                    Settings.embed_model = default_embedding
                    logger.info("✅ Modelo por defecto de LlamaIndex configurado")
                
                # Configurar LlamaIndex con Gemini
                if Gemini is not None and GeminiEmbedding is not None:
                    # Usar gemini-2.5-flash (rápido y moderno)
                    self._llm = Gemini(model="gemini-2.5-flash", api_key=api_key)
                    self._embedding_model = GeminiEmbedding(
                        model_name="models/embedding-001",
                        api_key=api_key
                    )
                    logger.info("✅ Modelos LlamaIndex configurados: gemini-2.5-flash + embedding-001")
                
                if Settings is not None and SimpleNodeParser is not None:
                    try:
                        # API moderna - configurar modelos explícitamente
                        Settings.llm = self._llm
                        Settings.embed_model = self._embedding_model
                        Settings.node_parser = SimpleNodeParser.from_defaults(
                            chunk_size=1024,
                            chunk_overlap=200
                        )
                        logger.info("✅ Settings de LlamaIndex configurados correctamente")
                    except Exception as e:
                        logger.error(f"❌ Error configurando Settings de LlamaIndex: {str(e)}")
                        raise
                
                logger.info("Modelos LlamaIndex inicializados correctamente")
            except Exception as e:
                logger.error(f"Error inicializando modelos LlamaIndex: {str(e)}")
                self._llm = None
                self._embedding_model = None
        else:
            logger.warning("GEMINI_API_KEY no configurado")
            self._llm = None
            self._embedding_model = None
            
        self._initialized = True
    
    async def load_tenant_documents(self, tenant_id: str, documentation_bucket_url: str) -> bool:
        """
        Carga documentos del bucket del tenant y crea un índice
        Usa persistencia en DB para evitar reprocesamiento
        
        Args:
            tenant_id: ID del tenant
            documentation_bucket_url: URL del bucket con documentos
            
        Returns:
            True si se cargaron exitosamente, False en caso contrario
        """
        try:
            if not LLAMA_INDEX_SUPPORT:
                logger.warning("LlamaIndex no disponible - cargando documentos sin indexación")
                # Cargar documentos sin indexación vectorial
                return await self._load_documents_simple(tenant_id, documentation_bucket_url)
            
            self._ensure_models_initialized()
            
            if not self._llm or not self._embedding_model:
                logger.warning("⚠️ Modelos de IA no disponibles (API key inválida/expirada) - cargando documentos en modo simple")
                # Fallback a modo simple cuando la API key falla
                return await self._load_documents_simple(tenant_id, documentation_bucket_url)
            
            logger.info(f"📚 Cargando documentos para tenant {tenant_id} desde: {documentation_bucket_url}")
            
            # 🚀 OPTIMIZACIÓN: Verificar si ya tenemos un índice en cache
            if tenant_id in self._index_cache:
                logger.info(f"✅ Índice ya existe en cache para tenant {tenant_id}")
                # 🔧 Asegurar que también tenemos documentos en modo simple como backup
                if tenant_id in self._document_cache:
                    doc_info = self._document_cache[tenant_id]
                    if 'documents' not in doc_info:
                        logger.info(f"📥 Cargando documentos en modo simple como backup...")
                        try:
                            await self._load_documents_simple(tenant_id, documentation_bucket_url)
                            logger.info(f"✅ Documentos en modo simple cargados como backup")
                        except Exception as e:
                            logger.warning(f"⚠️ No se pudieron cargar documentos en modo simple (no crítico): {e}")
                return True
            
            # 💾 CARGAR desde disco local si existe (rápido ~1s)
            import os
            storage_path = f"./storage/indices/{tenant_id}"
            if os.path.exists(storage_path) and os.listdir(storage_path):
                try:
                    logger.info(f"💾 Índice encontrado en disco: {storage_path}")
                    storage_context = StorageContext.from_defaults(persist_dir=storage_path)
                    index = load_index_from_storage(storage_context)
                    
                    self._index_cache[tenant_id] = index
                    self._document_cache[tenant_id] = {
                        "bucket_url": documentation_bucket_url,
                        "document_count": 0,
                        "loaded_from_disk": True
                    }
                    logger.info(f"✅ Índice cargado desde disco en ~1s para tenant {tenant_id}")
                    
                    # 🔧 IMPORTANTE: También cargar documentos en modo simple como backup
                    # para cuando el índice vectorial falle por API key
                    logger.info(f"📥 Cargando documentos en modo simple como backup...")
                    try:
                        await self._load_documents_simple(tenant_id, documentation_bucket_url)
                        logger.info(f"✅ Documentos en modo simple cargados como backup")
                    except Exception as simple_error:
                        logger.warning(f"⚠️ No se pudieron cargar documentos en modo simple (no crítico): {simple_error}")
                    
                    return True
                except Exception as load_error:
                    logger.warning(f"⚠️ Error cargando desde disco: {load_error}")
                    # Si falla cargar desde disco, intentar modo simple
                    logger.info(f"📥 Intentando cargar documentos en modo simple como fallback...")
                    return await self._load_documents_simple(tenant_id, documentation_bucket_url)
            
            # No hay persistencia, cargar desde bucket
            logger.info(f"📥 Cargando desde bucket (10-30s para primera vez)...")
            # 🗄️ NUEVO: Verificar si existe índice en DB (para saber si debería cargar)
            from chatbot_ai_service.services.document_index_persistence_service import document_index_persistence_service
            index_exists = document_index_persistence_service.index_exists(tenant_id, documentation_bucket_url)
            
            if index_exists:
                logger.info(f"🗄️ Metadatos encontrados en DB para tenant {tenant_id}")
                logger.info(f"📥 Cargando documentos desde bucket...")
            
            # Obtener lista de documentos del bucket
            logger.info(f"📥 Obteniendo lista de documentos desde bucket...")
            documents = await self._fetch_documents_from_bucket(documentation_bucket_url)
            
            if not documents:
                logger.warning(f"⚠️ No se encontraron documentos en el bucket para tenant {tenant_id} (URL: {documentation_bucket_url})")
                return False
            
            # Log de tipos de archivos encontrados
            file_types = {}
            for doc in documents:
                ext = doc['filename'].split('.')[-1].lower()
                file_types[ext] = file_types.get(ext, 0) + 1
            logger.info(f"📊 Documentos encontrados: {len(documents)} archivos - {dict(file_types)}")
            
            # Crear documentos de LlamaIndex
            logger.info(f"📖 Procesando {len(documents)} documentos...")
            
            # 🚀 OPTIMIZACIÓN LOCAL: Procesar TODOS los documentos incluso en desarrollo LOCAL
            import os
            # Detectar si estamos en desarrollo local usando variable específica
            is_local_dev = os.getenv("LOCAL_DEVELOPMENT", "false").lower() == "true"
            # Procesar TODOS los documentos disponibles para máximo contexto
            max_docs = len(documents)  # Sin límites - procesar todo
            documents_to_process = documents[:max_docs]
            
            logger.info(f"📚 Procesando {len(documents_to_process)} documentos (máximo {max_docs} - {'LOCAL DEV' if is_local_dev else 'REMOTE/PROD'})")
            
            llama_documents = []
            total_chars = 0
            for idx, doc_info in enumerate(documents_to_process, 1):
                try:
                    logger.info(f"📥 [{idx}/{len(documents_to_process)}] Descargando {doc_info['filename']}...")
                    content = await self._download_document_content(doc_info["url"])
                    if content:
                        document = Document(
                            text=content,
                            metadata={
                                "tenant_id": tenant_id,
                                "filename": doc_info["filename"],
                                "url": doc_info["url"],
                                "content_type": doc_info.get("content_type", "text/plain")
                            }
                        )
                        llama_documents.append(document)
                        total_chars += len(content)
                        logger.info(f"✅ [{idx}/{len(documents_to_process)}] {doc_info['filename']} - {len(content):,} chars")
                    else:
                        logger.warning(f"⚠️ [{idx}/{len(documents_to_process)}] {doc_info['filename']} - contenido vacío")
                except Exception as e:
                    logger.error(f"❌ Error cargando documento {doc_info['filename']}: {str(e)}")
                    continue
            
            if not llama_documents:
                logger.warning(f"⚠️ No se pudieron cargar documentos válidos para tenant {tenant_id}")
                return False
            
            # Crear índice vectorial
            logger.info(f"🔨 Creando índice vectorial con {len(llama_documents)} documentos...")
            index = VectorStoreIndex.from_documents(llama_documents)
            
            # 💾 GUARDAR índice para persistencia
            import os
            
            # Path local para desarrollo y para subir a GCS
            local_path = f"./storage/indices/{tenant_id}"
            
            # 🔧 Extraer nombre del bucket dinámicamente desde la URL
            # Formato: gs://bucket-name/... o https://storage.googleapis.com/bucket-name/...
            if documentation_bucket_url.startswith("gs://"):
                gcs_bucket = documentation_bucket_url.split("gs://")[1].split("/")[0]
            elif "storage.googleapis.com" in documentation_bucket_url:
                # https://storage.googleapis.com/bucket-name/...
                parts = documentation_bucket_url.split("storage.googleapis.com/")
                if len(parts) > 1:
                    gcs_bucket = parts[1].split("/")[0]
                else:
                    gcs_bucket = "unknown-bucket"  # Fallback
            else:
                # Intentar extraer bucket de cualquier formato
                import re
                match = re.search(r'gs://([^/]+)|storage\.googleapis\.com/([^/]+)', documentation_bucket_url)
                gcs_bucket = match.group(1) or match.group(2) if match else "unknown-bucket"
            
            logger.info(f"🔧 Usando bucket: {gcs_bucket} para tenant {tenant_id}")
            gcs_path = f"gs://{gcs_bucket}/indices/{tenant_id}"
            
            os.makedirs(local_path, exist_ok=True)
            try:
                # Guardar localmente
                index.storage_context.persist(persist_dir=local_path)
                logger.info(f"💾 Índice guardado localmente: {local_path}")
                
                # Opcional: Subir a Cloud Storage para persistencia en Cloud Run
                # (Puedes activar esto si necesitas persistencia entre reinicios)
                # await self._upload_index_to_gcs(local_path, gcs_path)
                
            except Exception as persist_error:
                logger.warning(f"⚠️ No se pudo persistir índice: {persist_error}")
            
            # Guardar en cache RAM
            print(f"🔍 DEBUG: Guardando en cache para tenant {tenant_id}")
            self._index_cache[tenant_id] = index
            self._document_cache[tenant_id] = {
                "bucket_url": documentation_bucket_url,
                "document_count": len(llama_documents),
                "total_chars": total_chars
            }
            
            print(f"🔍 DEBUG: Cache actualizado - Tenants en cache: {list(self._document_cache.keys())}")
            print(f"🔍 DEBUG: Index cache tiene keys: {list(self._index_cache.keys())}")
            
            # 🗄️ NUEVO: Guardar metadatos en DB para persistencia
            from chatbot_ai_service.services.document_index_persistence_service import document_index_persistence_service
            await document_index_persistence_service.save_index_metadata(
                tenant_id=tenant_id,
                bucket_url=documentation_bucket_url,
                documents_count=len(llama_documents),
                file_types=file_types,
                total_chars=total_chars
            )
            
            # Debug: Mostrar estado del cache después de guardar
            logger.info(f"🔍 DEBUG: Cache actualizado - Tenants: {list(self._document_cache.keys())}")
            
            logger.info(
                f"✅ Índice creado exitosamente para tenant {tenant_id} | "
                f"{len(llama_documents)} documentos | "
                f"{total_chars:,} caracteres totales | "
                f"~{total_chars // 1000}K tokens estimados"
            )
            logger.info(f"🗄️ Metadatos guardados en DB para tenant {tenant_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error cargando documentos para tenant {tenant_id}: {str(e)}")
            return False
    
    async def get_relevant_context(self, tenant_id: str, query: str, max_results: int = 3) -> str:
        """
        Obtiene contexto relevante para una consulta usando cache optimizado
        OPTIMIZADO: Reduce max_results por defecto para respuestas más rápidas
        
        Args:
            tenant_id: ID del tenant
            query: Consulta del usuario
            max_results: Número máximo de resultados a retornar (reducido por defecto)
            
        Returns:
            Contexto relevante como string
        """
        try:
            # Debug: Mostrar estado del cache
            logger.info(f"🔍 DEBUG: Cache actual - Tenants: {list(self._document_cache.keys())}")
            logger.info(f"🔍 DEBUG: Buscando tenant {tenant_id}")
            
            # Verificar si tenemos documentos en cache
            if tenant_id not in self._document_cache:
                logger.warning(f"⚠️ No hay documentos cargados para tenant {tenant_id}")
                logger.warning(f"🔍 DEBUG: Tenants disponibles en cache: {list(self._document_cache.keys())}")
                
                # Intentar cargar documentos obteniendo bucket_url desde la configuración del tenant
                try:
                    from chatbot_ai_service.services.configuration_service import configuration_service
                    tenant_config = configuration_service.get_tenant_config(tenant_id)
                    if tenant_config:
                        # Obtener bucket_url desde ai_config o documentación
                        ai_config = tenant_config.get('ai_config', {})
                        bucket_url = ai_config.get('documentation_bucket_url') or ai_config.get('documentationBucketUrl')
                        
                        if bucket_url:
                            logger.info(f"📥 Intentando cargar documentos en modo simple para tenant {tenant_id} desde {bucket_url[:50]}...")
                            loaded = await self._load_documents_simple(tenant_id, bucket_url)
                            if loaded:
                                logger.info(f"✅ Documentos cargados en modo simple para tenant {tenant_id}")
                            else:
                                logger.warning(f"⚠️ No se pudieron cargar documentos para tenant {tenant_id}")
                        else:
                            logger.debug(f"⚠️ No se encontró bucket_url en configuración del tenant {tenant_id}")
                except Exception as e:
                    logger.warning(f"⚠️ Error intentando cargar documentos: {e}")
                
                # Verificar nuevamente si ahora tenemos documentos
                if tenant_id not in self._document_cache:
                    return ""
            
            doc_info = self._document_cache[tenant_id]
            
            # 🚀 PRIORIDAD 1: Usar índice vectorial si está disponible (búsqueda semántica)
            print(f"🔍 DEBUG: ¿Tenant en index_cache? {tenant_id in self._index_cache}")
            print(f"🔍 DEBUG: ¿LLAMA_INDEX_SUPPORT? {LLAMA_INDEX_SUPPORT}")
            print(f"🔍 DEBUG: ¿VectorStoreIndex disponible? {VectorStoreIndex is not None}")
            print(f"🔍 DEBUG: Keys en _index_cache: {list(self._index_cache.keys())}")
            
            if tenant_id in self._index_cache and LLAMA_INDEX_SUPPORT and VectorStoreIndex is not None:
                print(f"🔍 DEBUG: ✅ USANDO ÍNDICE VECTORIAL para tenant {tenant_id}")
                logger.info(f"🔍 DEBUG: Usando índice vectorial para tenant {tenant_id}")
                index = self._index_cache[tenant_id]
                print(f"🔍 DEBUG: Creando query_engine...")
                query_engine = index.as_query_engine(
                    similarity_top_k=max_results,
                    response_mode="compact"
                )
                print(f"🔍 DEBUG: Haciendo query con: '{query[:100]}...'")
                try:
                    response = query_engine.query(query)
                    print(f"🔍 DEBUG: Respuesta obtenida de query_engine")
                    logger.info(f"🔍 DEBUG: Respuesta del índice vectorial: {response}")
                except Exception as query_error:
                    error_msg = str(query_error)
                    print(f"🔍 DEBUG: ERROR en query_engine.query: {query_error}")
                    logger.warning(f"⚠️ Error en índice vectorial (API key o embedding): {error_msg}")
                    
                    # Si es error de API key (403, expired, leaked, invalid), usar SmartRetriever como fallback
                    api_key_errors = ["403", "api key", "leaked", "expired", "invalid", "permission denied"]
                    if any(err in error_msg.lower() for err in api_key_errors):
                        logger.warning("⚠️ API key de Gemini inválida/expirada - usando SmartRetriever como fallback")
                        if tenant_id in self._document_cache:
                            doc_info = self._document_cache[tenant_id]
                            if 'documents' in doc_info:
                                from chatbot_ai_service.retrievers.smart_retriever import SmartRetriever
                                retriever = SmartRetriever()
                                search_results = retriever.search_documents(doc_info['documents'], query, max_results)
                                if search_results:
                                    context_parts = []
                                    for result in search_results:
                                        context_parts.append(f"[Documento] {result.filename}:\n{result.content[:500]}...")
                                    logger.info(f"✅ SmartRetriever encontró {len(search_results)} documentos relevantes")
                                    return "\n\n".join(context_parts)
                                else:
                                    logger.warning("⚠️ SmartRetriever no encontró documentos relevantes")
                            else:
                                logger.warning("⚠️ No hay documentos en modo simple para SmartRetriever")
                                # Intentar cargar documentos ahora
                                try:
                                    bucket_url = doc_info.get('bucket_url')
                                    if bucket_url:
                                        logger.info(f"📥 Cargando documentos en modo simple ahora...")
                                        await self._load_documents_simple(tenant_id, bucket_url)
                                        # Reintentar búsqueda
                                        if tenant_id in self._document_cache and 'documents' in self._document_cache[tenant_id]:
                                            retriever = SmartRetriever()
                                            search_results = retriever.search_documents(self._document_cache[tenant_id]['documents'], query, max_results)
                                            if search_results:
                                                context_parts = [f"[Documento] {r.filename}:\n{r.content[:500]}..." for r in search_results]
                                                return "\n\n".join(context_parts)
                                except Exception as load_err:
                                    logger.warning(f"⚠️ Error cargando documentos en modo simple: {load_err}")
                    
                    return ""
                
                # Extraer contexto relevante de la respuesta
                if hasattr(response, 'source_nodes'):
                    print(f"🔍 DEBUG: Encontrados {len(response.source_nodes)} nodos en la respuesta vectorial")
                    context_parts = []
                    for i, node in enumerate(response.source_nodes[:max_results]):
                        filename = node.metadata.get('filename', 'Documento')
                        score = getattr(node, 'score', 'N/A')
                        print(f"🔍 DEBUG: Nodo {i+1}: {filename} (score: {score})")
                        context_parts.append(f"**{filename}**:\n{node.text}")
                    
                    context = "\n\n".join(context_parts)
                    print(f"🔍 DEBUG: Contexto vectorial generado: {len(context)} caracteres, {len(context_parts)} nodos")
                    logger.info(f"🔍 DEBUG: Contexto vectorial generado: {len(context)} caracteres")
                    return context
                else:
                    logger.info(f"🔍 DEBUG: No hay source_nodes en la respuesta vectorial")
                    return ""
            else:
                print(f"🔍 DEBUG: ❌ NO HAY ÍNDICE VECTORIAL - usando SmartRetriever como fallback")
                logger.info(f"🔍 DEBUG: No hay índice vectorial, usando SmartRetriever")
                # Usar SmartRetriever como fallback
                from chatbot_ai_service.retrievers.smart_retriever import SmartRetriever
                retriever = SmartRetriever()
                
                logger.info(f"🔍 DEBUG: Buscando con SmartRetriever: '{query}'")
                search_results = retriever.search_documents(doc_info['documents'], query, max_results)
                logger.info(f"🔍 DEBUG: SmartRetriever encontró {len(search_results)} resultados")
                
                if search_results:
                    context_parts = []
                    for i, result in enumerate(search_results):
                        logger.info(f"🔍 DEBUG: Resultado {i+1}: {result.filename} (score: {result.score}, tipo: {result.match_type})")
                        context_parts.append(f"[Documento {i+1}] {result.content}")
                    
                    context = "\n\n".join(context_parts)
                    logger.info(f"🔍 DEBUG: Contexto generado: {len(context)} caracteres")
                    return context
                else:
                    logger.info(f"🔍 DEBUG: SmartRetriever no encontró resultados relevantes")
                    return ""
            
            # FALLBACK: Si no hay índice vectorial, usar búsqueda simple
            logger.warning(f"🔍 DEBUG: No hay índice vectorial disponible para tenant {tenant_id}")
            return ""
                
        except Exception as e:
            logger.error(f"Error obteniendo contexto para tenant {tenant_id}: {str(e)}")
            return ""
    
    async def _fetch_documents_from_bucket(self, bucket_url: str) -> List[Dict[str, Any]]:
        """
        Obtiene la lista de documentos disponibles en el bucket
        
        Args:
            bucket_url: URL del bucket
            
        Returns:
            Lista de documentos con metadatos
        """
        try:
            # Parsear URL del bucket
            parsed_url = urlparse(bucket_url)
            
            if "googleapis.com" in parsed_url.netloc:
                # Google Cloud Storage
                return await self._fetch_gcs_documents(bucket_url)
            elif "s3" in parsed_url.netloc or "amazonaws.com" in parsed_url.netloc:
                # Amazon S3
                return await self._fetch_s3_documents(bucket_url)
            else:
                # Asumir que es una URL directa a un documento o lista
                return await self._fetch_direct_documents(bucket_url)
                
        except Exception as e:
            logger.error(f"Error obteniendo documentos del bucket {bucket_url}: {str(e)}")
            return []
    
    async def _fetch_gcs_documents(self, bucket_url: str) -> List[Dict[str, Any]]:
        """Obtiene documentos de Google Cloud Storage"""
        try:
            # Parsear URL de GCS
            # Formato: https://storage.googleapis.com/bucket-name
            parsed_url = urlparse(bucket_url)
            # Para URLs como https://storage.googleapis.com/bucket-name
            if parsed_url.netloc == "storage.googleapis.com":
                bucket_name = parsed_url.path.lstrip("/")
            else:
                bucket_name = parsed_url.netloc
            
            if not bucket_name:
                logger.error(f"No se pudo extraer nombre del bucket de: {bucket_url}")
                return []
            
            logger.info(f"Obteniendo documentos del bucket GCS: {bucket_name}")
            
            # Usar la API pública de GCS para listar objetos
            list_url = f"https://storage.googleapis.com/storage/v1/b/{bucket_name}/o"
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(list_url)
                
                if response.status_code == 200:
                    data = response.json()
                    documents = []
                    
                    for item in data.get("items", []):
                        # Filtrar solo archivos de texto, PDF, Word y Excel
                        name = item.get("name", "")
                        if name.endswith(('.txt', '.md', '.pdf', '.docx', '.doc', '.xlsx', '.xls')):
                            # Construir URL pública del archivo
                            file_url = f"https://storage.googleapis.com/storage/v1/b/{bucket_name}/o/{name}?alt=media"
                            
                            documents.append({
                                "filename": name,
                                "url": file_url,
                                "content_type": item.get("contentType", "text/plain"),
                                "size": item.get("size", "0")
                            })
                    
                    logger.info(f"Encontrados {len(documents)} documentos en bucket GCS: {bucket_name}")
                    return documents
                else:
                    logger.error(f"Error accediendo al bucket GCS {bucket_name}: {response.status_code}")
                    return []
                    
        except Exception as e:
            logger.error(f"Error obteniendo documentos de GCS: {str(e)}")
            return []
    
    async def _fetch_s3_documents(self, bucket_url: str) -> List[Dict[str, Any]]:
        """Obtiene documentos de Amazon S3"""
        try:
            # Para S3, necesitaríamos usar boto3
            # Por ahora, retornamos una lista vacía
            logger.warning("Soporte para Amazon S3 no implementado aún")
            return []
        except Exception as e:
            logger.error(f"Error obteniendo documentos de S3: {str(e)}")
            return []
    
    async def _fetch_direct_documents(self, url: str) -> List[Dict[str, Any]]:
        """Obtiene documentos de una URL directa"""
        try:
            # Si es una URL directa a un documento
            if url.endswith(('.pdf', '.txt', '.docx', '.md', '.xlsx', '.xls')):
                return [{
                    "filename": url.split('/')[-1],
                    "url": url,
                    "content_type": self._get_content_type(url)
                }]
            
            # Si es una URL a una lista o directorio, intentar obtener contenido
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                
                if response.status_code == 200:
                    # Intentar parsear como JSON si es una API
                    try:
                        data = response.json()
                        if isinstance(data, list):
                            return data
                        elif isinstance(data, dict) and "documents" in data:
                            return data["documents"]
                    except:
                        # Si no es JSON, asumir que es texto plano
                        pass
                
                logger.warning(f"No se pudo obtener lista de documentos de {url}")
                return []
                
        except Exception as e:
            logger.error(f"Error obteniendo documentos directos de {url}: {str(e)}")
            return []
    
    async def _download_document_content(self, url: str) -> Optional[str]:
        """
        Descarga el contenido de un documento
        
        Args:
            url: URL del documento
            
        Returns:
            Contenido del documento como string
        """
        try:
            logger.debug(f"🌐 Descargando desde: {url}")
            async with httpx.AsyncClient(timeout=10.0) as client:
                response = await client.get(url)
                
                if response.status_code == 200:
                    content_type = response.headers.get('content-type', '')
                    
                    if 'text/' in content_type or url.endswith(('.txt', '.md')):
                        return response.text
                    elif 'application/pdf' in content_type or url.endswith('.pdf'):
                        if PDF_SUPPORT:
                            return self._extract_pdf_text(response.content)
                        else:
                            logger.warning("Soporte para PDF no disponible - PyPDF2 no instalado")
                            return None
                    # IMPORTANTE: Chequear Excel ANTES de Word (Excel también contiene 'openxmlformats')
                    elif 'spreadsheetml.sheet' in content_type or url.endswith(('.xlsx', '.xls')):
                        if EXCEL_SUPPORT:
                            return self._extract_excel_text(response.content)
                        else:
                            logger.warning("Soporte para Excel no disponible - openpyxl no instalado")
                            return None
                    elif 'wordprocessingml.document' in content_type or url.endswith('.docx'):
                        if DOCX_SUPPORT:
                            return self._extract_docx_text(response.content)
                        else:
                            logger.warning("Soporte para Word no disponible - python-docx no instalado")
                            return None
                    else:
                        # Intentar como texto plano
                        try:
                            return response.text
                        except:
                            logger.warning(f"No se pudo procesar documento: {url}")
                            return None
                else:
                    logger.error(f"Error descargando documento {url}: {response.status_code}")
                    return None
                    
        except Exception as e:
            logger.error(f"Error descargando documento {url}: {str(e)}")
            return None
    
    def _get_content_type(self, url: str) -> str:
        """Determina el tipo de contenido basado en la extensión del archivo"""
        extension = url.lower().split('.')[-1]
        
        content_types = {
            'txt': 'text/plain',
            'md': 'text/markdown',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'xls': 'application/vnd.ms-excel',
            'rtf': 'application/rtf'
        }
        
        return content_types.get(extension, 'text/plain')
    
    def clear_tenant_cache(self, tenant_id: str):
        """Limpia el cache de documentos para un tenant específico"""
        self._index_cache.pop(tenant_id, None)
        self._document_cache.pop(tenant_id, None)
        logger.info(f"Cache limpiado para tenant: {tenant_id}")
    
    def get_tenant_document_info(self, tenant_id: str) -> Optional[Dict[str, Any]]:
        """Obtiene información sobre los documentos cargados para un tenant"""
        return self._document_cache.get(tenant_id)
    
    async def _load_documents_simple(self, tenant_id: str, documentation_bucket_url: str) -> bool:
        """Carga documentos de forma simple sin indexación vectorial"""
        try:
            logger.info(f"Cargando documentos simple para tenant {tenant_id} desde: {documentation_bucket_url}")
            
            # Obtener lista de documentos del bucket
            documents = await self._fetch_documents_from_bucket(documentation_bucket_url)
            
            if not documents:
                logger.warning(f"No se encontraron documentos en el bucket para tenant {tenant_id}")
                return False
            
            # Procesar documentos y guardar contenido
            document_contents = []
            for doc_info in documents:
                try:
                    content = await self._download_document_content(doc_info["url"])
                    if content:
                        document_contents.append({
                            "filename": doc_info["filename"],
                            "content": content,
                            "url": doc_info["url"]
                        })
                        logger.info(f"Documento cargado: {doc_info['filename']}")
                except Exception as e:
                    logger.error(f"Error cargando documento {doc_info['filename']}: {str(e)}")
                    continue
            
            if not document_contents:
                logger.warning(f"No se pudieron cargar documentos válidos para tenant {tenant_id}")
                return False
            
            # Guardar en cache simple
            self._document_cache[tenant_id] = {
                "bucket_url": documentation_bucket_url,
                "document_count": len(document_contents),
                "documents": document_contents,
                "loaded_at": "simple_mode"
            }
            
            logger.info(f"Documentos cargados en modo simple para tenant {tenant_id}: {len(document_contents)} documentos")
            return True
            
        except Exception as e:
            logger.error(f"Error cargando documentos simple para tenant {tenant_id}: {str(e)}")
            return False
    
    def _extract_pdf_text(self, pdf_content: bytes) -> Optional[str]:
        """Extrae texto de un archivo PDF"""
        try:
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            extracted_text = "\n".join(text_content)
            logger.info(f"Texto extraído de PDF: {len(extracted_text)} caracteres")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de PDF: {str(e)}")
            return None
    
    def _extract_docx_text(self, docx_content: bytes) -> Optional[str]:
        """Extrae texto de un archivo Word"""
        try:
            docx_file = io.BytesIO(docx_content)
            doc = DocxDocument(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            extracted_text = "\n".join(text_content)
            logger.info(f"Texto extraído de Word: {len(extracted_text)} caracteres")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de Word: {str(e)}")
            return None
    
    def _extract_excel_text(self, excel_content: bytes) -> Optional[str]:
        """Extrae texto de un archivo Excel"""
        try:
            excel_file = io.BytesIO(excel_content)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text_content = []
            
            # Procesar todas las hojas
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Agregar nombre de la hoja como sección
                text_content.append(f"\n=== {sheet_name} ===\n")
                
                # Extraer todas las celdas con contenido
                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            # Convertir el valor a string
                            value = str(cell.value).strip()
                            if value:
                                row_values.append(value)
                    
                    # Solo agregar filas con contenido
                    if row_values:
                        text_content.append(" | ".join(row_values))
            
            extracted_text = "\n".join(text_content)
            logger.info(f"Texto extraído de Excel: {len(extracted_text)} caracteres de {len(workbook.sheetnames)} hoja(s)")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de Excel: {str(e)}")
            return None
    
    def _simple_search(self, documents: List[Dict[str, Any]], query: str, max_results: int) -> str:
        """Búsqueda optimizada usando SmartRetriever"""
        try:
            from chatbot_ai_service.retrievers.smart_retriever import smart_retriever
            
            # Usar el nuevo SmartRetriever
            results = smart_retriever.search_documents(documents, query, max_results)
            
            if not results:
                logger.info("🔍 No se encontraron documentos relevantes")
                return ""
            
            # Construir contexto usando el SmartRetriever
            context = smart_retriever.get_context_from_results(results, max_context_length=2000)
            
            logger.info(f"🔍 Búsqueda optimizada: {len(results)} documentos encontrados")
            return context
            
        except Exception as e:
            logger.error(f"Error en búsqueda optimizada: {str(e)}")
            # Fallback a búsqueda básica si hay error
            return self._basic_search_fallback(documents, query, max_results)
    
    def _basic_search_fallback(self, documents: List[Dict[str, Any]], query: str, max_results: int) -> str:
        """Búsqueda básica de fallback en caso de error"""
        try:
            query_lower = query.lower()
            scored_documents = []
            
            for doc in documents:
                content = doc["content"].lower()
                filename = doc.get("filename", "").lower()
                score = 0
                
                # Búsqueda simple por palabras
                if query_lower in content:
                    score += 100
                if query_lower in filename:
                    score += 150
                
                # Búsqueda por palabras individuales
                for word in query_lower.split():
                    if len(word) > 2:
                        score += content.count(word)
                        if word in filename:
                            score += 20
                
                if score > 0:
                    scored_documents.append((score, doc))
            
            # Ordenar y tomar los mejores
            scored_documents.sort(key=lambda x: x[0], reverse=True)
            top_documents = scored_documents[:max_results]
            
            # Construir contexto básico
            context_parts = []
            for score, doc in top_documents:
                content_preview = doc["content"][:500] + "..." if len(doc["content"]) > 500 else doc["content"]
                context_parts.append(f"**{doc['filename']}**:\n{content_preview}")
            
            return "\n\n".join(context_parts)
            
        except Exception as e:
            logger.error(f"Error en búsqueda de fallback: {str(e)}")
            return ""


# Instancia global
document_context_service = DocumentContextService()
